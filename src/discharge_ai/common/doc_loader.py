"""
common/doc_loader.py
====================

Reads every incoming clinical document format the hospital network produces and
returns plain text plus (when the source is structured) the parsed payload:

    .txt / .md      → read as UTF-8
    .json           → parsed + rendered to readable text
    .pdf            → pdfplumber, then pypdf, then the `.ocr.txt` sidecar
    .png/.jpg/.jpeg → the `.ocr.txt` sidecar, else Tesseract if installed
    .docx           → python-docx

Scanned / handwritten samples in `Data/incoming` ship a pre-extracted
`<name>.ocr.txt` sidecar (e.g. `P1022_daan_bakker.png.ocr.txt`).  The sidecar is
always preferred over live OCR: it is what the hospital's scanning pipeline
produced, and it keeps the demo reproducible on machines without Tesseract.

Also owns patient-file *discovery*: mapping a patient id to its discharge
report, lab report and bill, which the MCP Clinical Watcher tool exposes.
"""

from __future__ import annotations

import json
import logging
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from ..settings import settings
from .schemas import DocType

log = logging.getLogger(__name__)

TEXT_SUFFIXES = {".txt", ".md", ".text"}
JSON_SUFFIXES = {".json"}
PDF_SUFFIXES = {".pdf"}
IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp"}
DOCX_SUFFIXES = {".docx"}

SUPPORTED_SUFFIXES = (
    TEXT_SUFFIXES | JSON_SUFFIXES | PDF_SUFFIXES | IMAGE_SUFFIXES | DOCX_SUFFIXES
)

#  Sidecars are companions, never documents in their own right.
OCR_SIDECAR_MARKER = ".ocr.txt"

#  NB: `\b` would not fire on "P1019_bill.json" because `_` is a word
#  character, so the boundaries are spelled out explicitly.
PATIENT_ID_RE = re.compile(r"(?<![A-Za-z0-9])(P\d{4})(?![0-9])", re.IGNORECASE)


# --------------------------------------------------------------------------- #
#  Loaded document
# --------------------------------------------------------------------------- #
@dataclass
class LoadedDocument:
    path: Path
    doc_type: str
    text: str = ""
    data: dict[str, Any] | None = None       # populated for JSON sources
    fmt: str = "txt"
    ocr_used: bool = False
    tables: list[list[list[str]]] = field(default_factory=list)
    error: str | None = None

    @property
    def name(self) -> str:
        return self.path.name

    def to_payload(self) -> dict[str, Any]:
        return {
            "file": self.name,
            "path": str(self.path),
            "doc_type": self.doc_type,
            "format": self.fmt,
            "ocr_used": self.ocr_used,
            "characters": len(self.text),
            "has_structured_data": self.data is not None,
            "error": self.error,
        }


# --------------------------------------------------------------------------- #
#  Format readers
# --------------------------------------------------------------------------- #
def _read_text_file(path: Path) -> str:
    for encoding in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_bytes().decode("utf-8", errors="replace")


def _sidecar_for(path: Path) -> Path | None:
    """Locate the pre-extracted OCR text that ships next to a scan/PDF.

    Handles both naming styles present in the dataset:
        P1020_labs.pdf            → P1020_labs.pdf.ocr.txt
        P1022_daan_bakker.png.jpeg → P1022_daan_bakker.png.ocr.txt
    """
    candidates = [
        path.with_name(path.name + OCR_SIDECAR_MARKER),           # foo.pdf.ocr.txt
        path.with_suffix(path.suffix + OCR_SIDECAR_MARKER),       # same, explicit
        path.with_name(path.stem + OCR_SIDECAR_MARKER),           # foo.png.ocr.txt
        path.with_name(Path(path.stem).stem + OCR_SIDECAR_MARKER),  # foo.ocr.txt
    ]
    for candidate in candidates:
        if candidate.exists():
            return candidate

    #  Real scanning pipelines mis-key filenames: the dataset ships
    #  `P1023_grace_benett.png.jpeg` next to `P1023_grace_bennett.png.ocr.txt`
    #  (one 'n' short).  Fall back to the patient id, which is unambiguous.
    patient_id = extract_patient_id(path)
    if patient_id:
        for sibling in sorted(path.parent.glob(f"*{OCR_SIDECAR_MARKER}")):
            if extract_patient_id(sibling) == patient_id:
                log.info(
                    "Matched OCR sidecar %s to %s by patient id (filenames differ)",
                    sibling.name, path.name,
                )
                return sibling
    return None


def _read_pdf(path: Path) -> tuple[str, list[list[list[str]]], bool]:
    """(text, tables, ocr_used) — pdfplumber → pypdf → OCR sidecar."""
    tables: list[list[list[str]]] = []

    try:
        import pdfplumber

        pages, extracted_tables = [], []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                pages.append(page.extract_text() or "")
                for table in page.extract_tables() or []:
                    extracted_tables.append(
                        [[(cell or "") for cell in row] for row in table]
                    )
        text = "\n".join(pages).strip()
        tables = extracted_tables
        if text:
            return text, tables, False
    except Exception as exc:  # noqa: BLE001
        log.debug("pdfplumber failed on %s: %s", path.name, exc)

    try:
        from pypdf import PdfReader

        text = "\n".join((p.extract_text() or "") for p in PdfReader(str(path)).pages)
        if text.strip():
            return text.strip(), tables, False
    except Exception as exc:  # noqa: BLE001
        log.debug("pypdf failed on %s: %s", path.name, exc)

    sidecar = _sidecar_for(path)
    if sidecar:
        log.info("Using OCR sidecar %s for %s", sidecar.name, path.name)
        return _read_text_file(sidecar).strip(), tables, True

    return "", tables, False


def _read_image(path: Path) -> tuple[str, bool]:
    """(text, ocr_used) — sidecar first, live Tesseract only as a fallback."""
    sidecar = _sidecar_for(path)
    if sidecar:
        return _read_text_file(sidecar).strip(), True

    try:
        import pytesseract
        from PIL import Image

        log.info("Running Tesseract OCR on %s", path.name)
        return pytesseract.image_to_string(Image.open(path)).strip(), True
    except Exception as exc:  # noqa: BLE001
        log.warning(
            "No OCR text for %s (%s). Add a '<name>%s' sidecar or install "
            "Tesseract (see requirements-optional.txt).",
            path.name, exc, OCR_SIDECAR_MARKER,
        )
        return "", False


def _read_docx(path: Path) -> str:
    try:
        import docx

        document = docx.Document(str(path))
        parts = [para.text for para in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text.strip() for cell in row.cells))
        return "\n".join(parts).strip()
    except Exception as exc:  # noqa: BLE001
        log.warning("docx read failed for %s: %s", path.name, exc)
        return ""


def json_to_text(data: Any, indent: int = 0) -> str:
    """Flatten a JSON payload into readable "Label: value" lines.

    Keeps the original (possibly non-English) keys — the Normalizer Agent
    translates afterwards — and keeps `raw_text` blocks intact, which is where
    the Hindi/Dutch source narratives live in this dataset.
    """
    pad = "  " * indent
    lines: list[str] = []

    if isinstance(data, dict):
        for key, value in data.items():
            label = str(key).replace("_", " ").title() if key.isascii() else str(key)
            if isinstance(value, (dict, list)):
                nested = json_to_text(value, indent + 1)
                if nested.strip():
                    lines.append(f"{pad}{label}:")
                    lines.append(nested)
            elif value is not None and str(value) != "":
                lines.append(f"{pad}{label}: {value}")
    elif isinstance(data, list):
        for index, item in enumerate(data, start=1):
            if isinstance(item, (dict, list)):
                nested = json_to_text(item, indent + 1)
                if nested.strip():
                    lines.append(f"{pad}- [{index}]")
                    lines.append(nested)
            elif item is not None:
                lines.append(f"{pad}- {item}")
    else:
        lines.append(f"{pad}{data}")

    return "\n".join(lines)


# --------------------------------------------------------------------------- #
#  Public loader
# --------------------------------------------------------------------------- #
def load_document(path: str | Path, doc_type: str = "unknown") -> LoadedDocument:
    """Read one clinical document of any supported format."""
    path = Path(path)
    doc = LoadedDocument(path=path, doc_type=doc_type)

    if not path.exists():
        doc.error = f"File not found: {path}"
        return doc

    suffix = path.suffix.lower()
    try:
        if suffix in JSON_SUFFIXES:
            doc.fmt = "json"
            raw = _read_text_file(path)
            payload = json.loads(raw)
            doc.data = payload if isinstance(payload, dict) else {"items": payload}
            doc.text = json_to_text(payload)

        elif suffix in TEXT_SUFFIXES:
            doc.fmt = "txt"
            doc.text = _read_text_file(path).strip()

        elif suffix in PDF_SUFFIXES:
            doc.fmt = "pdf"
            doc.text, doc.tables, doc.ocr_used = _read_pdf(path)

        elif suffix in IMAGE_SUFFIXES:
            doc.fmt = "image"
            doc.text, doc.ocr_used = _read_image(path)

        elif suffix in DOCX_SUFFIXES:
            doc.fmt = "docx"
            doc.text = _read_docx(path)

        else:
            doc.fmt = suffix.lstrip(".") or "unknown"
            doc.error = f"Unsupported file format: {suffix}"
            return doc

        if not doc.text.strip() and not doc.data:
            doc.error = "No text could be extracted"

    except Exception as exc:  # noqa: BLE001
        doc.error = f"{type(exc).__name__}: {exc}"
        log.warning("Failed to load %s: %s", path.name, doc.error)

    return doc


# --------------------------------------------------------------------------- #
#  Patient-file discovery  (backs the MCP Clinical Watcher tool)
# --------------------------------------------------------------------------- #
def _is_sidecar(path: Path) -> bool:
    return path.name.endswith(OCR_SIDECAR_MARKER)


def _doc_type_for_dir(directory: Path) -> str:
    name = directory.name.lower()
    if "doctor" in name or "discharge" in name or "report" in name and "lab" not in name:
        return DocType.DISCHARGE_REPORT.value
    if "lab" in name:
        return DocType.LAB_REPORT.value
    if "bill" in name:
        return DocType.BILL.value
    return "unknown"


def _folder_map() -> dict[str, Path]:
    return {
        DocType.DISCHARGE_REPORT.value: settings.path("doctor_reports"),
        DocType.LAB_REPORT.value: settings.path("lab_reports"),
        DocType.BILL.value: settings.path("bills"),
    }


def extract_patient_id(path: Path) -> str | None:
    match = PATIENT_ID_RE.search(path.name)
    return match.group(1).upper() if match else None


def scan_incoming(input_root: Path | None = None) -> dict[str, dict[str, list[str]]]:
    """Discover `{patient_id: {doc_type: [file, ...]}}` under the input root."""
    root = Path(input_root) if input_root else settings.path("input_root")
    catalogue: dict[str, dict[str, list[str]]] = {}

    for doc_type, folder in _folder_map().items():
        if not folder.exists():
            continue
        for path in sorted(folder.iterdir()):
            if not path.is_file() or _is_sidecar(path):
                continue
            if path.suffix.lower() not in SUPPORTED_SUFFIXES:
                continue
            patient_id = extract_patient_id(path)
            if not patient_id:
                continue
            try:
                relative = str(path.relative_to(root))
            except ValueError:
                relative = str(path)
            catalogue.setdefault(patient_id, {}).setdefault(doc_type, []).append(relative)

    return dict(sorted(catalogue.items()))


def list_patient_ids() -> list[str]:
    return list(scan_incoming().keys())


def find_patient_files(patient_id: str) -> dict[str, Path]:
    """`{doc_type: Path}` for one patient (first match per document type)."""
    patient_id = patient_id.upper()
    found: dict[str, Path] = {}

    for doc_type, folder in _folder_map().items():
        if not folder.exists():
            continue
        for path in sorted(folder.iterdir()):
            if not path.is_file() or _is_sidecar(path):
                continue
            if path.suffix.lower() not in SUPPORTED_SUFFIXES:
                continue
            if extract_patient_id(path) == patient_id:
                # Prefer structured/textual sources over scans when both exist.
                current = found.get(doc_type)
                if current is None or _format_rank(path) < _format_rank(current):
                    found[doc_type] = path

    return found


def _format_rank(path: Path) -> int:
    """Lower is better: structured data beats text beats PDF beats scan."""
    suffix = path.suffix.lower()
    if suffix in JSON_SUFFIXES:
        return 0
    if suffix in TEXT_SUFFIXES:
        return 1
    if suffix in DOCX_SUFFIXES:
        return 2
    if suffix in PDF_SUFFIXES:
        return 3
    return 4


def load_patient_documents(patient_id: str) -> dict[str, LoadedDocument]:
    """Load all three document types for one patient."""
    return {
        doc_type: load_document(path, doc_type)
        for doc_type, path in find_patient_files(patient_id).items()
    }


# --------------------------------------------------------------------------- #
#  MCP Roots path-traversal guard
# --------------------------------------------------------------------------- #
class RootAccessError(PermissionError):
    """Raised when a request tries to escape the declared MCP Root."""


def resolve_within_root(candidate: str | Path, root: str | Path) -> Path:
    """Resolve `candidate` and prove it lives inside `root`.

    This is the `Path.relative_to()` check the specification requires of the
    Primary MCP Server: the Clinical Watcher tool discovers folders through
    `ctx.list_roots()`, and any path that resolves outside the declared root —
    `../../etc/passwd`, an absolute path, a symlink escape — is rejected.
    """
    root_path = Path(root).resolve()
    target = Path(candidate)
    target = (root_path / target).resolve() if not target.is_absolute() else target.resolve()

    try:
        target.relative_to(root_path)
    except ValueError as exc:
        raise RootAccessError(
            f"Access denied: {target} is outside the declared MCP root {root_path}"
        ) from exc

    return target
